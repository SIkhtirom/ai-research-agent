"""Word (.docx) text extraction using python-docx, including OCR of embedded images."""

import logging
from pathlib import Path

from docx import Document

from .base_parser import DocumentParser, ExtractedDocument
from .ocr import extract_images_from_zip, is_ocr_available, ocr_image_bytes

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    def parse(self, source: Path | str) -> ExtractedDocument:
        source_path = Path(source)
        document = Document(source_path)

        paragraph_texts = (paragraph.text for paragraph in document.paragraphs)
        body_text = "\n".join(text for text in paragraph_texts if text.strip())

        # Extract text embedded in pictures/diagrams inside the document.
        image_ocr_texts: list[str] = []
        for data in extract_images_from_zip(source_path):
            ocr_text = ocr_image_bytes(data)
            if ocr_text:
                image_ocr_texts.append(f"(Gambar: {ocr_text})")

        sections: list[str] = []
        if body_text:
            sections.append(body_text)
        if image_ocr_texts:
            sections.append("\n".join(image_ocr_texts))

        content_text = "\n\n".join(sections)
        if not content_text.strip():
            logger.warning(
                "DOCX '%s' produced no extractable text (it may be entirely image-based).",
                source_path.name,
            )

        metadata = {
            "filename": source_path.name,
            "paragraph_count": len(document.paragraphs),
            "image_count": len(image_ocr_texts),
            "ocr_available": is_ocr_available(),
            "author": document.core_properties.author,
            "title": document.core_properties.title,
        }
        return ExtractedDocument(content_text=content_text, metadata=metadata)
